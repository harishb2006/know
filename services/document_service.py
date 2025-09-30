import os
import uuid
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
import json

import numpy as np
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import io

from app.database import get_db_sync
from app.models import Document, DocumentChunk
from app.config import settings
from .ocr_service import OCRService
from .embedding_service import EmbeddingService

logger = logging.getLogger(__name__)

class DocumentService:
    """Enhanced document processing with OCR and embeddings"""
    
    def __init__(self):
        self.ocr_service = OCRService()
        self.embedding_service = EmbeddingService()
        
    def upload_document(self, file_data: bytes, filename: str, is_public: bool = False) -> Dict[str, Any]:
        """Upload and process document with full text extraction and embedding generation"""
        try:
            # Determine file type
            file_extension = filename.lower().split('.')[-1]
            
            # Extract text content
            extraction_result = self._extract_text_content(file_data, file_extension)
            
            if not extraction_result['success']:
                return {
                    'success': False,
                    'error': f"Failed to extract text: {extraction_result.get('error', 'Unknown error')}"
                }
            
            # Save file to disk
            file_id = str(uuid.uuid4())
            file_path = os.path.join(settings.upload_dir, f"{file_id}_{filename}")
            
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            # Create document record
            db = get_db_sync()
            try:
                document = Document(
                    id=file_id,
                    filename=filename,
                    file_type=file_extension,
                    content=extraction_result['text'],
                    file_size=len(file_data),
                    is_public=is_public,
                    metadata={
                        'file_path': file_path,
                        'extraction_method': extraction_result.get('method', 'unknown'),
                        'extraction_confidence': extraction_result.get('confidence', 0),
                        'ocr_details': extraction_result.get('ocr_details', {}),
                        'total_pages': extraction_result.get('total_pages', 1)
                    }
                )
                
                db.add(document)
                db.commit()
                
                # Create chunks and embeddings
                chunks_created = self._create_document_chunks(db, document)
                
                logger.info(f"Document uploaded: {filename} ({chunks_created} chunks created)")
                
                return {
                    'success': True,
                    'document_id': file_id,
                    'filename': filename,
                    'chunks_created': chunks_created,
                    'text_preview': extraction_result['text'][:500] + "..." if len(extraction_result['text']) > 500 else extraction_result['text'],
                    'metadata': document.metadata
                }
                
            except Exception as e:
                db.rollback()
                raise e
            finally:
                db.close()
                
        except Exception as e:
            logger.error(f"Document upload failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_text_content(self, file_data: bytes, file_extension: str) -> Dict[str, Any]:
        """Extract text from various file formats"""
        
        if file_extension == 'pdf':
            return self._extract_from_pdf(file_data)
        elif file_extension in ['docx', 'doc']:
            return self._extract_from_docx(file_data)
        elif file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']:
            return self._extract_from_image(file_data)
        elif file_extension == 'txt':
            return self._extract_from_text(file_data)
        else:
            return {
                'success': False,
                'error': f"Unsupported file type: {file_extension}"
            }
    
    def _extract_from_pdf(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            pdf_reader = PdfReader(io.BytesIO(file_data))
            text_content = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                text_content += f"\n--- Page {page_num + 1} ---\n"
                text_content += page.extract_text()
            
            return {
                'success': True,
                'text': text_content.strip(),
                'method': 'pdf_extraction',
                'total_pages': len(pdf_reader.pages),
                'confidence': 95  # PDF text extraction is usually reliable
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_from_docx(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            doc = DocxDocument(io.BytesIO(file_data))
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return {
                'success': True,
                'text': text_content.strip(),
                'method': 'docx_extraction',
                'total_pages': 1,
                'confidence': 98  # DOCX extraction is very reliable
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_from_image(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        ocr_result = self.ocr_service.extract_text_from_image(file_data)
        
        if ocr_result['success']:
            return {
                'success': True,
                'text': ocr_result['text'],
                'method': 'ocr_extraction',
                'confidence': ocr_result['extraction_confidence'],
                'ocr_details': {
                    'words_detected': len(ocr_result['words']),
                    'word_confidences': [w['confidence'] for w in ocr_result['words']]
                }
            }
        else:
            return {
                'success': False,
                'error': ocr_result.get('error', 'OCR failed')
            }
    
    def _extract_from_text(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            text_content = file_data.decode('utf-8')
            
            return {
                'success': True,
                'text': text_content,
                'method': 'text_file',
                'confidence': 100
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _create_document_chunks(self, db, document: Document) -> int:
        """Split document into chunks and create embeddings"""
        try:
            text = document.content
            chunks = self._split_text_into_chunks(text)
            
            chunks_created = 0
            for i, chunk_text in enumerate(chunks):
                # Generate embedding
                embedding = self.embedding_service.generate_embedding(chunk_text)
                
                # Calculate character positions
                start_char = i * (settings.chunk_size - settings.chunk_overlap)
                end_char = min(start_char + len(chunk_text), len(text))
                
                # Create chunk record
                chunk = DocumentChunk(
                    document_id=document.id,
                    content=chunk_text,
                    chunk_index=i,
                    start_char=start_char,
                    end_char=end_char,
                    embedding=embedding.tolist() if embedding is not None else None
                )
                
                db.add(chunk)
                chunks_created += 1
            
            db.commit()
            return chunks_created
            
        except Exception as e:
            logger.error(f"Chunk creation failed: {e}")
            db.rollback()
            return 0
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= settings.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + settings.chunk_size
            
            # If not the last chunk, try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence boundary
                last_period = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)
                last_space = text.rfind(' ', start, end)
                
                break_point = max(last_period, last_newline, last_space)
                if break_point > start:
                    end = break_point + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - settings.chunk_overlap
        
        return chunks
    
    def search_documents(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """Search documents using vector similarity"""
        try:
            # Generate query embedding
            query_embedding = self.embedding_service.generate_embedding(query)
            
            if query_embedding is None:
                return []
            
            # Get all chunks with embeddings
            db = get_db_sync()
            try:
                chunks = db.query(DocumentChunk).filter(
                    DocumentChunk.embedding.isnot(None)
                ).all()
                
                if not chunks:
                    return []
                
                # Calculate similarities
                similarities = []
                for chunk in chunks:
                    chunk_embedding = np.array(chunk.embedding)
                    similarity = self._cosine_similarity(query_embedding, chunk_embedding)
                    
                    similarities.append({
                        'chunk': chunk,
                        'similarity': similarity
                    })
                
                # Sort by similarity and return top k
                similarities.sort(key=lambda x: x['similarity'], reverse=True)
                
                results = []
                for item in similarities[:k]:
                    chunk = item['chunk']
                    document = chunk.document
                    
                    results.append({
                        'document_id': document.id,
                        'document_name': document.filename,
                        'chunk_content': chunk.content,
                        'similarity_score': float(item['similarity']),
                        'chunk_index': chunk.chunk_index,
                        'start_char': chunk.start_char,
                        'end_char': chunk.end_char,
                        'page_number': chunk.page_number,
                        'is_public': document.is_public
                    })
                
                return results
                
            finally:
                db.close()
                
        except Exception as e:
            logger.error(f"Document search failed: {e}")
            return []
    
    def _cosine_similarity(self, a: np.ndarray, b: np.ndarray) -> float:
        """Calculate cosine similarity between two vectors"""
        dot_product = np.dot(a, b)
        norm_a = np.linalg.norm(a)
        norm_b = np.linalg.norm(b)
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        
        return dot_product / (norm_a * norm_b)
    
    def get_public_documents(self, skip: int = 0, limit: int = 50) -> List[Dict[str, Any]]:
        """Get list of public documents"""
        db = get_db_sync()
        try:
            documents = db.query(Document).filter(
                Document.is_public == True
            ).offset(skip).limit(limit).all()
            
            return [
                {
                    'id': doc.id,
                    'filename': doc.filename,
                    'file_type': doc.file_type,
                    'file_size': doc.file_size,
                    'created_at': doc.created_at.isoformat(),
                    'metadata': doc.metadata
                }
                for doc in documents
            ]
            
        finally:
            db.close()
    
    def get_document_by_id(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get document details by ID"""
        db = get_db_sync()
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            
            if not document:
                return None
            
            return {
                'id': document.id,
                'filename': document.filename,
                'file_type': document.file_type,
                'content': document.content,
                'file_size': document.file_size,
                'is_public': document.is_public,
                'metadata': document.metadata,
                'created_at': document.created_at.isoformat(),
                'updated_at': document.updated_at.isoformat()
            }
            
        finally:
            db.close()